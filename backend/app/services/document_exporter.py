#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档导出服务
支持将解析结果导出为 Word、Markdown 等格式
"""

import logging
from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime

logger = logging.getLogger(__name__)


class DocumentExporter:
    """文档导出器"""
    
    def __init__(self):
        self.export_dir = Path("exports")
        self.export_dir.mkdir(parents=True, exist_ok=True)
        logger.info("文档导出服务已初始化")
    
    def export_to_markdown(self, analysis_result: Dict[str, Any], filename: str = None) -> Dict[str, Any]:
        """将解析结果导出为 Markdown 文档"""
        try:
            # 生成文件名
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"document_export_{timestamp}.md"
            
            file_path = self.export_dir / filename
            
            # 构建 Markdown 内容
            markdown_content = self._build_markdown_content(analysis_result)
            
            # 写入文件
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            logger.info(f"Markdown 文档导出成功: {file_path}")
            
            return {
                "success": True,
                "file_path": str(file_path),
                "file_name": filename,
                "file_size": file_path.stat().st_size,
                "format": "markdown"
            }
            
        except Exception as e:
            logger.error(f"Markdown 导出失败: {e}")
            return {
                "success": False,
                "error": str(e),
                "format": "markdown"
            }
    
    def export_to_word(self, analysis_result: Dict[str, Any], filename: str = None) -> Dict[str, Any]:
        """将解析结果导出为 Word 文档"""
        try:
            # 检查 python-docx 是否可用
            try:
                from docx import Document
                from docx.shared import Pt, RGBColor
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                DOCX_AVAILABLE = True
            except ImportError:
                DOCX_AVAILABLE = False
                logger.warning("python-docx 不可用，请安装: pip install python-docx")
                return {
                    "success": False,
                    "error": "python-docx 未安装，请运行: pip install python-docx",
                    "format": "word"
                }
            
            # 生成文件名
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"document_export_{timestamp}.docx"
            
            file_path = self.export_dir / filename
            
            # 创建 Word 文档
            doc = Document()
            
            # 添加标题
            title = doc.add_heading('文档分析报告', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 添加文档信息
            doc.add_heading('文档信息', level=1)
            info_table = doc.add_table(rows=2, cols=2)
            info_table.style = 'Light Grid Accent 1'
            
            info_table.rows[0].cells[0].text = '文件路径'
            info_table.rows[0].cells[1].text = analysis_result.get('file_path', 'N/A')
            info_table.rows[1].cells[0].text = '文件类型'
            info_table.rows[1].cells[1].text = analysis_result.get('file_type', 'N/A')
            
            # 添加提取文本
            if analysis_result.get('extracted_text'):
                doc.add_heading('提取文本', level=1)
                doc.add_paragraph(analysis_result['extracted_text'])
            
            # 添加化学实体
            if analysis_result.get('chemical_entities') and len(analysis_result['chemical_entities']) > 0:
                doc.add_heading('化学实体', level=1)
                
                # 创建表格
                entity_table = doc.add_table(rows=1, cols=4)
                entity_table.style = 'Light Grid Accent 1'
                
                # 设置表头
                headers = ['实体名称', '类型', '置信度', '位置']
                for i, header in enumerate(headers):
                    cell = entity_table.rows[0].cells[i]
                    cell.text = header
                    cell.paragraphs[0].runs[0].bold = True
                
                # 添加数据
                for entity in analysis_result['chemical_entities']:
                    row_cells = entity_table.add_row().cells
                    row_cells[0].text = entity.get('text', 'N/A')
                    row_cells[1].text = entity.get('type', 'N/A')
                    row_cells[2].text = f"{entity.get('confidence', 0) * 100:.1f}%"
                    pos = entity.get('position', {})
                    row_cells[3].text = f"{pos.get('start', 0)}-{pos.get('end', 0)}"
            
            # 添加工艺参数
            if analysis_result.get('process_parameters') and len(analysis_result['process_parameters']) > 0:
                doc.add_heading('工艺参数', level=1)
                
                # 创建表格
                param_table = doc.add_table(rows=1, cols=4)
                param_table.style = 'Light Grid Accent 1'
                
                # 设置表头
                headers = ['参数名称', '数值', '单位', '置信度']
                for i, header in enumerate(headers):
                    cell = param_table.rows[0].cells[i]
                    cell.text = header
                    cell.paragraphs[0].runs[0].bold = True
                
                # 添加数据
                for param in analysis_result['process_parameters']:
                    row_cells = param_table.add_row().cells
                    row_cells[0].text = param.get('name', 'N/A')
                    row_cells[1].text = param.get('value', 'N/A')
                    row_cells[2].text = param.get('unit', 'N/A')
                    row_cells[3].text = f"{param.get('confidence', 0) * 100:.1f}%"
            
            # 添加元数据
            if analysis_result.get('metadata'):
                doc.add_heading('元数据', level=1)
                for key, value in analysis_result['metadata'].items():
                    doc.add_paragraph(f"{key}: {value}")
            
            # 添加导出时间
            doc.add_paragraph()
            doc.add_paragraph(f'导出时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            
            # 保存文档
            doc.save(file_path)
            
            logger.info(f"Word 文档导出成功: {file_path}")
            
            return {
                "success": True,
                "file_path": str(file_path),
                "file_name": filename,
                "file_size": file_path.stat().st_size,
                "format": "word"
            }
            
        except Exception as e:
            logger.error(f"Word 导出失败: {e}")
            return {
                "success": False,
                "error": str(e),
                "format": "word"
            }
    
    def _build_markdown_content(self, analysis_result: Dict[str, Any]) -> str:
        """构建 Markdown 内容"""
        lines = []
        
        # 标题
        lines.append("# 文档分析报告")
        lines.append("")
        
        # 文档信息
        lines.append("## 文档信息")
        lines.append("")
        lines.append(f"- **文件路径**: {analysis_result.get('file_path', 'N/A')}")
        lines.append(f"- **文件类型**: {analysis_result.get('file_type', 'N/A')}")
        lines.append(f"- **解析状态**: {'✅ 成功' if analysis_result.get('success') else '❌ 失败'}")
        if analysis_result.get('error'):
            lines.append(f"- **错误信息**: {analysis_result['error']}")
        lines.append("")
        
        # 提取文本
        if analysis_result.get('extracted_text'):
            lines.append("## 提取文本")
            lines.append("")
            text = analysis_result['extracted_text']
            # 将文本按段落分割
            paragraphs = text.split('\n')
            for para in paragraphs:
                if para.strip():
                    lines.append(para.strip())
            lines.append("")
        
        # 化学实体
        if analysis_result.get('chemical_entities') and len(analysis_result['chemical_entities']) > 0:
            lines.append("## 化学实体")
            lines.append("")
            lines.append("| 实体名称 | 类型 | 置信度 | 位置 |")
            lines.append("|---------|------|--------|------|")
            
            for entity in analysis_result['chemical_entities']:
                name = entity.get('text', 'N/A')
                entity_type = entity.get('type', 'N/A')
                confidence = f"{entity.get('confidence', 0) * 100:.1f}%"
                pos = entity.get('position', {})
                position = f"{pos.get('start', 0)}-{pos.get('end', 0)}"
                lines.append(f"| {name} | {entity_type} | {confidence} | {position} |")
            
            lines.append("")
        
        # 工艺参数
        if analysis_result.get('process_parameters') and len(analysis_result['process_parameters']) > 0:
            lines.append("## 工艺参数")
            lines.append("")
            lines.append("| 参数名称 | 数值 | 单位 | 置信度 |")
            lines.append("|---------|------|------|--------|")
            
            for param in analysis_result['process_parameters']:
                name = param.get('name', 'N/A')
                value = param.get('value', 'N/A')
                unit = param.get('unit', 'N/A')
                confidence = f"{param.get('confidence', 0) * 100:.1f}%"
                lines.append(f"| {name} | {value} | {unit} | {confidence} |")
            
            lines.append("")
        
        # 元数据
        if analysis_result.get('metadata'):
            lines.append("## 元数据")
            lines.append("")
            for key, value in analysis_result['metadata'].items():
                lines.append(f"- **{key}**: {value}")
            lines.append("")
        
        # 导出时间
        lines.append(f"---")
        lines.append("")
        lines.append(f"*导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*")
        lines.append("")
        
        return "\n".join(lines)
    
    def get_export_formats(self) -> Dict[str, Any]:
        """获取支持的导出格式"""
        return {
            "formats": [
                {
                    "name": "Markdown",
                    "extension": ".md",
                    "description": "Markdown 格式，支持文本、表格和基本格式",
                    "available": True
                },
                {
                    "name": "Word",
                    "extension": ".docx",
                    "description": "Microsoft Word 文档，支持丰富的格式和样式",
                    "available": self._check_docx_available()
                }
            ],
            "export_directory": str(self.export_dir)
        }
    
    def _check_docx_available(self) -> bool:
        """检查 python-docx 是否可用"""
        try:
            import docx
            return True
        except ImportError:
            return False