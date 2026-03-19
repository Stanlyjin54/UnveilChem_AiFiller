#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档解析器 - 使用python-docx
"""

import os
import io
from pathlib import Path
from typing import Dict, List, Any, Optional
import logging

try:
    from docx import Document
    from docx.shared import Inches
    import pandas as pd
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("Word解析依赖缺失,请安装python-docx依赖")

# OCR依赖
try:
    import pytesseract
    from PIL import Image, ImageEnhance
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    Image = None
    ImageEnhance = None
    logging.warning("OCR依赖缺失,请安装pillow, pytesseract")

from . import BaseDocumentParser

class WordParser(BaseDocumentParser):
    """Word文档解析器"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.docx', '.doc']
        self.parser_name = "WORD_PARSER_V1"
    
    def _enhance_image_for_ocr(self, image) -> object:
        """增强图像以提高OCR识别率"""
        if ImageEnhance is None:
            return image
        
        # 转换为灰度
        if image.mode != 'L':
            image = image.convert('L')
        
        # 对比度增强
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(1.5)
        
        # 降噪处理
        from PIL import ImageFilter
        image = image.filter(ImageFilter.MedianFilter())
        
        return image
    
    def can_parse(self, file_path: Path) -> bool:
        """检查是否是Word文件"""
        return file_path.suffix.lower() in ['.docx', '.doc']
    
    def parse(self, file_path: Path) -> Dict[str, Any]:
        """解析Word文件"""
        if not DOCX_AVAILABLE:
            return {
                "success": False,
                "error": "Word解析功能不可用，请安装python-docx依赖",
                "file_path": str(file_path),
                "parser_used": self.parser_name,
                "metadata": self.get_metadata(file_path)
            }
        
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        result = {
            "success": False,
            "parser_used": self.parser_name,
            "file_path": str(file_path),
            "metadata": self.get_metadata(file_path),
            "text_content": "",
            "tables": [],
            "images": [],
            "parameters": {},
            "errors": []
        }
        
        try:
            doc = Document(str(file_path))
            
            # 提取文本内容
            all_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    all_text.append(para.text.strip())
            
            # 提取图片并进行OCR识别
            image_texts = []
            image_count = 0
            
            # 遍历文档中的所有部分
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    # 提取图片
                    image_count += 1
                    image_part = rel.target_part
                    image_bytes = image_part.blob
                    
                    # 进行OCR识别
                    if OCR_AVAILABLE and Image is not None:
                        try:
                            # 将图片字节转换为PIL Image
                            image = Image.open(io.BytesIO(image_bytes))
                            
                            # 增强图像以提高OCR识别率
                            enhanced_image = self._enhance_image_for_ocr(image)
                            
                            ocr_text = pytesseract.image_to_string(enhanced_image, lang='chi_sim+eng')
                            if ocr_text.strip():
                                image_texts.append(f"=== 图片{image_count} (OCR) ===\n{ocr_text}")
                                result["images"].append(f"image_{image_count}.png")
                        except Exception as ocr_error:
                            result["errors"].append(f"图片{image_count} OCR识别失败: {str(ocr_error)}")
                            logging.warning(f"图片{image_count} OCR识别失败: {ocr_error}")
                    else:
                        result["images"].append(f"image_{image_count}.png")
            
            # 合并文本内容和图片OCR结果
            if image_texts:
                all_text.extend(image_texts)
            
            result["text_content"] = "\n".join(all_text)
            
            # 提取表格
            tables = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                
                if table_data:
                    tables.append({
                        "table_index": table_idx,
                        "data": table_data,
                        "rows": len(table_data),
                        "cols": len(table_data[0]) if table_data else 0
                    })
            
            result["tables"] = tables
            
            # 提取参数
            if result["text_content"]:
                result["parameters"] = self._extract_parameters(result["text_content"])
            
            result["success"] = True
            result["metadata"]["tables_found"] = len(tables)
            result["metadata"]["paragraphs_count"] = len([p for p in doc.paragraphs if p.text.strip()])
            result["metadata"]["images_found"] = image_count
            result["metadata"]["images_ocr_processed"] = len(image_texts)
            
        except Exception as e:
            result["errors"].append(f"Word解析失败: {str(e)}")
            logging.error(f"Word解析错误: {e}")
        
        return result
    
    def _extract_parameters(self, text: str) -> Dict[str, Any]:
        """提取工艺参数"""
        import re
        parameters = {}
        
        # 温度参数
        temp_patterns = [
            r'(?:温度|Temperature|TEMP)[:\s]*([\d.]+)\s*[°℃]?C?',
            r'([\d.]+)\s*[°℃](?:\s*温度|\s*Temperature)?',
            r'T\s*[=]\s*([\d.]+)\s*[°℃]?'
        ]
        
        # 压力参数
        pressure_patterns = [
            r'(?:压力|Pressure|PRESS)[:\s]*([\d.]+)\s*(?:MPa|kPa|Pa)?',
            r'([\d.]+)\s*(?:MPa|kPa|Pa)(?:\s*压力|\s*Pressure)?',
            r'P\s*[=]\s*([\d.]+)\s*(?:MPa|kPa|Pa)?'
        ]
        
        # 流量参数
        flow_patterns = [
            r'(?:流量|Flow|FLOW)[:\s]*([\d.]+)\s*(?:m³/h|L/min|kg/h)?',
            r'([\d.]+)\s*(?:m³/h|L/min|kg/h)(?:\s*流量|\s*Flow)?'
        ]
        
        # 提取各类参数
        for pattern_list, param_type in [
            (temp_patterns, "temperature"),
            (pressure_patterns, "pressure"),
            (flow_patterns, "flow")
        ]:
            values = []
            for pattern in pattern_list:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    try:
                        value = float(match.group(1))
                        if self._validate_parameter_value(param_type, value):
                            values.append({
                                "value": value,
                                "unit": self._extract_unit(match.group(0), param_type),
                                "context": match.group(0).strip(),
                                "position": match.span()
                            })
                    except (ValueError, IndexError):
                        continue
            
            if values:
                parameters[param_type] = values
        
        return parameters
    
    def _validate_parameter_value(self, param_type: str, value: float) -> bool:
        """验证参数值是否合理"""
        validation_rules = {
            "temperature": (-100, 1000),  # -100°C 到 1000°C
            "pressure": (0, 50),          # 0 到 50 MPa
            "flow": (0, 10000)            # 0 到 10000 m³/h
        }
        
        if param_type in validation_rules:
            min_val, max_val = validation_rules[param_type]
            return min_val <= value <= max_val
        
        return True
    
    def _extract_unit(self, text: str, param_type: str) -> str:
        """提取参数单位"""
        import re
        
        unit_patterns = {
            "temperature": [r'°C', r'℃', r'C(?![a-zA-Z])'],
            "pressure": [r'MPa', r'kPa', r'Pa'],
            "flow": [r'm³/h', r'L/min', r'kg/h']
        }
        
        if param_type in unit_patterns:
            for pattern in unit_patterns[param_type]:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    return match.group(0)
        
        # 默认单位
        default_units = {
            "temperature": "°C",
            "pressure": "MPa", 
            "flow": "m³/h"
        }
        
        return default_units.get(param_type, "")