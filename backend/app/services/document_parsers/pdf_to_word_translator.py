#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PDF → Word → Ollama 翻译服务
"""

import asyncio
import aiohttp
import os
from pdf2docx import Converter
from docx import Document
from typing import Optional, Callable
import logging

logger = logging.getLogger(__name__)


class PDFToWordTranslator:
    """PDF 转 Word 并翻译"""
    
    def __init__(self, ollama_model="demonbyron/HY-MT1.5-1.8B:latest"):
        self.model = ollama_model
        self.base_url = "http://localhost:11434"
    
    async def translate_pdf_to_word(
        self,
        pdf_path: str,
        output_path: str,
        lang_in: str = "en",
        lang_out: str = "zh",
        progress_callback: Optional[Callable[[float], None]] = None
    ) -> str:
        """PDF → Word → 翻译"""
        
        logger.info(f"开始翻译: {pdf_path}")
        
        # 步骤 1: PDF → docx
        if progress_callback:
            progress_callback(0.1)
        
        temp_docx = pdf_path.replace('.pdf', '_temp.docx')
        logger.info(f"PDF → docx: {temp_docx}")
        
        try:
            cv = Converter(pdf_path)
            cv.convert(temp_docx)
            cv.close()
            
            logger.info("PDF → docx 转换完成")
        except Exception as e:
            logger.error(f"PDF → docx 转换失败: {e}")
            raise
        
        # 步骤 2: 读取 Word 文档
        if progress_callback:
            progress_callback(0.2)
        
        try:
            doc = Document(temp_docx)
            logger.info(f"Word 文档包含 {len(doc.paragraphs)} 个段落")
        except Exception as e:
            logger.error(f"读取 Word 文档失败: {e}")
            raise
        
        # 步骤 3: 翻译段落
        total_paragraphs = len(doc.paragraphs)
        translated_count = 0
        
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                try:
                    translated = await self._translate_text(
                        para.text,
                        lang_in,
                        lang_out
                    )
                    para.text = translated
                    translated_count += 1
                except Exception as e:
                    logger.error(f"翻译段落 {i+1} 失败: {e}")
                    para.text = f"[翻译失败: {para.text}]"
            
            if progress_callback and i % 5 == 0:
                progress_callback(0.2 + 0.7 * (i + 1) / total_paragraphs)
        
        logger.info(f"翻译完成: {translated_count} 个段落")
        
        # 步骤 4: 保存翻译后的 Word
        if progress_callback:
            progress_callback(0.95)
        
        try:
            doc.save(output_path)
            logger.info(f"保存翻译结果: {output_path}")
        except Exception as e:
            logger.error(f"保存 Word 文档失败: {e}")
            raise
        
        # 清理临时文件
        try:
            os.remove(temp_docx)
            logger.info(f"清理临时文件: {temp_docx}")
        except Exception as e:
            logger.warning(f"清理临时文件失败: {e}")
        
        if progress_callback:
            progress_callback(1.0)
        
        return output_path
    
    async def _translate_text(self, text: str, lang_in: str, lang_out: str) -> str:
        """使用 Ollama API 翻译"""
        async with aiohttp.ClientSession() as session:
            async with session.post(
                f"{self.base_url}/api/generate",
                json={
                    "model": self.model,
                    "prompt": f"translate {lang_in} to {lang_out}:{text}",
                    "stream": False
                },
                timeout=aiohttp.ClientTimeout(total=30)
            ) as resp:
                result = await resp.json()
                return result["response"]